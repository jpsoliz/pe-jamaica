from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = Path(__file__).resolve().parent
ENV_ROOT = (
    REPO_ROOT.parent
    / "python-envs"
    / "arcgispro-survey-ai"
)

DOCX_PATH = OUTPUT_DIR / "ArcGISPro37_Environment_Installation_Guide.docx"
CONDA_REQUIREMENTS_PATH = OUTPUT_DIR / "requirements-conda.txt"
PIP_REQUIREMENTS_PATH = OUTPUT_DIR / "requirements-pip.txt"
COMPONENT_INVENTORY_PATH = OUTPUT_DIR / "arcgispro-survey-ai-component-inventory.csv"


def read_conda_packages(env_root: Path) -> list[dict[str, str]]:
    packages: list[dict[str, str]] = []
    for json_path in sorted((env_root / "conda-meta").glob("*.json")):
        data = json.loads(json_path.read_text(encoding="utf-8"))
        packages.append(
            {
                "manager": "conda",
                "name": data.get("name", ""),
                "version": data.get("version", ""),
                "build": data.get("build", ""),
                "channel": data.get("channel", ""),
                "source": json_path.name,
                "installer": "conda",
                "role": "inventory",
            }
        )
    return packages


def normalize_name(name: str) -> str:
    return re.sub(r"[-_.]+", "-", name).lower()


def read_dist_info_name_version(dist_info: Path) -> tuple[str, str]:
    metadata = dist_info / "METADATA"
    name = ""
    version = ""
    if metadata.exists():
        for line in metadata.read_text(encoding="utf-8", errors="replace").splitlines():
            if line.startswith("Name: "):
                name = line[6:].strip()
            elif line.startswith("Version: "):
                version = line[9:].strip()
            if name and version:
                break
    return name, version


def read_pip_packages(env_root: Path, conda_names: set[str]) -> list[dict[str, str]]:
    pip_packages: list[dict[str, str]] = []
    site_packages = env_root / "Lib" / "site-packages"
    for dist_info in sorted(site_packages.glob("*.dist-info")):
        name, version = read_dist_info_name_version(dist_info)
        installer_path = dist_info / "INSTALLER"
        installer = installer_path.read_text(encoding="utf-8", errors="replace").strip() if installer_path.exists() else ""
        if installer != "pip" or not name:
            continue
        in_conda = normalize_name(name) in conda_names
        pip_packages.append(
            {
                "manager": "pip",
                "name": name,
                "version": version,
                "build": "",
                "channel": "pypi",
                "source": dist_info.name,
                "installer": installer,
                "role": "pip-overlay" if in_conda else "pip-only",
            }
        )
    return pip_packages


def extract_history_specs(env_root: Path) -> list[str]:
    history = env_root / "conda-meta" / "history"
    specs: list[str] = []
    if not history.exists():
        return specs
    current_cmd = ""
    for line in history.read_text(encoding="utf-8", errors="replace").splitlines():
        if line.startswith("# cmd:"):
            current_cmd = line
            continue
        if "# update specs:" not in line and "# install specs:" not in line:
            continue
        if " create --clone " in current_cmd:
            continue
        if " remove " in current_cmd:
            continue
        for match in re.findall(r"'([^']+)'", line):
            if "::" in match or "==" in match:
                continue
            if match not in specs:
                specs.append(match)
    return specs


def write_component_files(conda_packages: list[dict[str, str]], pip_packages: list[dict[str, str]], history_specs: list[str]) -> None:
    conda_specs = [spec for spec in history_specs if spec not in {"openssl", "certifi"}]
    if not conda_specs:
        conda_specs = ["openai"]

    CONDA_REQUIREMENTS_PATH.write_text("\n".join(conda_specs) + "\n", encoding="utf-8")
    PIP_REQUIREMENTS_PATH.write_text(
        "\n".join(f"{pkg['name']}=={pkg['version']}" for pkg in pip_packages) + "\n",
        encoding="utf-8",
    )

    with COMPONENT_INVENTORY_PATH.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=["manager", "name", "version", "build", "channel", "installer", "role", "source"],
        )
        writer.writeheader()
        for row in sorted(conda_packages + pip_packages, key=lambda r: (r["manager"], r["name"].lower())):
            writer.writerow(row)


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr_cells[index].text = header
        for paragraph in hdr_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_doc(conda_packages: list[dict[str, str]], pip_packages: list[dict[str, str]], history_specs: list[str]) -> None:
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ArcGIS Pro 3.7 Environment Installation Guide")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("AI Survey tool environment migration and installer preparation").italic = True

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This document defines the recommended process for preparing the Python environment used by the AI Survey "
        "tool on target computers running ArcGIS Pro 3.7. The current source environment was inspected from:"
    )
    add_code_block(doc, [str(ENV_ROOT)])

    doc.add_heading("Key Recommendation", level=1)
    add_bullets(
        doc,
        [
            "Do not copy the ArcGIS Pro 3.6 environment folder directly to ArcGIS Pro 3.7 computers.",
            "Do not modify the target computer's default arcgispro-py3 environment.",
            "Clone the clean ArcGIS Pro 3.7 default environment, then install only the additional packages required by the tool.",
            "Install conda packages first and pip packages last.",
            "Keep the final package lists in source control so the later installer can call the same steps.",
        ],
    )

    doc.add_heading("Extracted Components", level=1)
    doc.add_paragraph(
        f"The source environment contains {len(conda_packages)} conda package metadata records. Most of these are base "
        "ArcGIS Pro packages or transitive dependencies from the cloned default environment and should not be replayed "
        "blindly against ArcGIS Pro 3.7."
    )
    doc.add_paragraph(
        "The conda history indicates the only explicitly changed conda package after the clone was:"
    )
    add_table(doc, ["Manager", "Package", "Install guidance"], [["conda", "openai", "Install into the ArcGIS Pro 3.7 clone if not already present or if tool tests require it."]])

    doc.add_paragraph(
        f"The environment also has {len(pip_packages)} pip-installed distributions. These are captured in requirements-pip.txt "
        "and should be installed after conda packages."
    )
    add_table(
        doc,
        ["Package", "Version", "Role"],
        [[pkg["name"], pkg["version"], pkg["role"]] for pkg in pip_packages],
    )

    doc.add_heading("Repository Files Created", level=1)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["requirements-conda.txt", "Conda packages to install after cloning the ArcGIS Pro 3.7 default environment."],
            ["requirements-pip.txt", "Pip packages and exact versions observed in the source environment."],
            ["arcgispro-survey-ai-component-inventory.csv", "Full extracted inventory for auditing and installer design."],
            ["ArcGISPro37_Environment_Installation_Guide.docx", "This installation guide."],
        ],
    )

    doc.add_heading("Manual Installation Procedure", level=1)
    add_numbered(
        doc,
        [
            "Install ArcGIS Pro 3.7 on the target computer.",
            "Open the ArcGIS Pro Python Command Prompt as the user that will run the tool.",
            "Clone the clean default environment to a separate named environment.",
            "Activate the cloned environment.",
            "Install conda package requirements.",
            "Install pip package requirements.",
            "Run the verification commands.",
            "Configure the tool or add-in to use the cloned environment path.",
        ],
    )

    doc.add_heading("Command Sequence", level=1)
    add_code_block(
        doc,
        [
            "conda create --clone arcgispro-py3 --name arcgispro-survey-ai --pinned",
            "conda activate arcgispro-survey-ai",
            "conda install --file requirements-conda.txt -c esri -c conda-forge -c defaults",
            "pip install -r requirements-pip.txt",
        ],
    )

    doc.add_heading("Verification", level=1)
    doc.add_paragraph("Run these checks after package installation:")
    add_code_block(
        doc,
        [
            "python -c \"import sys; print(sys.executable)\"",
            "python -c \"import arcpy; print('arcpy OK')\"",
            "python -c \"import openai; import flask; import pdfplumber; import pypdfium2; print('AI Survey env OK')\"",
            "conda list > installed-conda-list.txt",
            "pip freeze > installed-pip-freeze.txt",
        ],
    )

    doc.add_heading("Installer Notes", level=1)
    add_bullets(
        doc,
        [
            "The installer should detect ArcGIS Pro 3.7 before creating the environment.",
            "The installer should create or update a cloned environment, not the default arcgispro-py3 environment.",
            "The installer should run conda before pip.",
            "The installer should keep logs from conda install, pip install, and verification imports.",
            "The installer should fail clearly if ArcGIS Pro 3.7 is missing or if arcpy cannot import from the cloned environment.",
        ],
    )

    doc.add_heading("Open Item Before Final Installer", level=1)
    doc.add_paragraph(
        "Before packaging the installer, compare this source environment against a clean ArcGIS Pro 3.7 default environment. "
        "Anything already present in ArcGIS Pro 3.7 should remain provided by Esri's base environment rather than being "
        "forced from the old environment. This reduces version conflicts during the ArcGIS Pro 3.6 to 3.7 migration."
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    if not ENV_ROOT.exists():
        raise SystemExit(f"Environment root not found: {ENV_ROOT}")
    conda_packages = read_conda_packages(ENV_ROOT)
    conda_names = {normalize_name(pkg["name"]) for pkg in conda_packages}
    pip_packages = read_pip_packages(ENV_ROOT, conda_names)
    history_specs = extract_history_specs(ENV_ROOT)
    write_component_files(conda_packages, pip_packages, history_specs)
    build_doc(conda_packages, pip_packages, history_specs)
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {CONDA_REQUIREMENTS_PATH}")
    print(f"Wrote {PIP_REQUIREMENTS_PATH}")
    print(f"Wrote {COMPONENT_INVENTORY_PATH}")


if __name__ == "__main__":
    main()
